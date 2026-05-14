"""
文档解析模块 - 支持多种文档格式
"""
import os
from pathlib import Path
from typing import List
import markdown
from docx import Document as DocxDocument
from PyPDF2 import PdfReader


class DocumentParser:
    """文档解析器 - 支持 docx, pdf, txt, md 格式"""

    @staticmethod
    def parse(file_path: str) -> str:
        """根据文件扩展名解析文档"""
        path = Path(file_path)
        ext = path.suffix.lower()

        parsers = {
            '.txt': DocumentParser._parse_txt,
            '.md': DocumentParser._parse_md,
            '.docx': DocumentParser._parse_docx,
            '.pdf': DocumentParser._parse_pdf,
        }

        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"不支持的文件格式: {ext}")

        return parser(file_path)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """解析 TXT 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _parse_md(file_path: str) -> str:
        """解析 Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
            html = markdown.markdown(md_content)
            text = DocumentParser._html_to_text(html)
            return text

    @staticmethod
    def _html_to_text(html: str) -> str:
        """简单的 HTML 转文本"""
        import re
        text = re.sub(r'<[^>]+>', ' ', html)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """解析 Word 文档"""
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append(' | '.join(cells))

        return '\n'.join(paragraphs)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """解析 PDF 文件"""
        reader = PdfReader(file_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"[第{page_num + 1}页]\n{text}")

        return '\n'.join(text_parts)

    @staticmethod
    def get_supported_formats() -> List[str]:
        """获取支持的文档格式"""
        return ['.txt', '.md', '.docx', '.pdf']
